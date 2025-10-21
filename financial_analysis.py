#!/usr/bin/env python3
"""
Módulo de Análisis Financiero para SuperBincent
Integra análisis contable automático con el sistema de impuestos
"""

import pandas as pd
import matplotlib.pyplot as plt
import os
import glob
import re
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import numpy as np
import logging
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from pathlib import Path

# Set pandas option to avoid FutureWarning
pd.set_option('future.no_silent_downcasting', True)

# Install openpyxl if missing
try:
    import openpyxl
except ImportError:
    print("Instalando openpyxl...")
    os.system("pip install openpyxl")
    import openpyxl

class FinancialAnalyzer:
    """Analizador financiero integrado con SuperBincent"""
    
    def __init__(self, downloads_dir: str = None):
        self.downloads_dir = downloads_dir or os.path.expanduser("~/Downloads")
        self.reports_dir = os.path.join(self.downloads_dir, "superbincent_reports")
        self.logger = logging.getLogger(__name__)
        
        # Crear directorio de reportes si no existe
        os.makedirs(self.reports_dir, exist_ok=True)
        
        # Configuración de email
        self.email_config = {
            'smtp_server': 'smtp.gmail.com',
            'smtp_port': 587,
            'sender_email': 'arielsanroj@carmanfe.com.co',
            'sender_password': 'jxxu siui wmfp dklj',
            'recipient_email': 'arielsanroj@carmanfe.com.co'
        }
        
        # Meses para ordenamiento
        self.month_order = ['ENERO', 'FEBRERO', 'MARZO', 'ABRIL', 'MAYO', 'JUNIO', 
                           'JULIO', 'AGOSTO', 'SEPTIEMBRE', 'OCTUBRE', 'NOVIEMBRE', 'DICIEMBRE']
        self.month_index = {m: i for i, m in enumerate(self.month_order)}
        
        # Datos financieros consolidados
        self.consolidated_data = {}
        self.current_month = None
        self.budget_data = {
            'ingresos_mensual': 100_000_000,
            'gastos_mensual': 125_000_000
        }
        
        self.logger.info("🧮 Analizador financiero inicializado")
    
    def detect_latest_financial_file(self) -> Optional[str]:
        """Detectar el archivo financiero más reciente"""
        pattern = os.path.join(self.downloads_dir, "INFORME DE * APRU- 2025 .xls")
        files = glob.glob(pattern)
        
        if not files:
            self.logger.warning(f"No se encontró archivo financiero con patrón: {pattern}")
            return None
        
        latest_file = max(files, key=os.path.getmtime)
        self.logger.info(f"📄 Archivo financiero detectado: {latest_file}")
        return latest_file
    
    def extract_month_from_filename(self, filename: str) -> str:
        """Extraer mes del nombre del archivo"""
        match = re.search(r'DE (\w+) APRU', filename)
        if not match:
            raise ValueError("No se pudo extraer el mes del nombre del archivo")
        return match.group(1).upper()
    
    def load_financial_data(self, file_path: str) -> Dict:
        """Cargar datos financieros del archivo Excel"""
        try:
            xls = pd.ExcelFile(file_path)
            self.current_month = self.extract_month_from_filename(os.path.basename(file_path))
            
            self.logger.info(f"📊 Cargando datos financieros para {self.current_month}")
            self.logger.info(f"Hojas disponibles: {xls.sheet_names}")
            
            # Cargar hojas principales
            data = {
                'balance_sheet_name': f'BALANCE {self.current_month}',
                'sheets': xls.sheet_names,
                'current_month': self.current_month
            }
            
            # Verificar que existe la hoja de balance
            if data['balance_sheet_name'] not in xls.sheet_names:
                raise ValueError(f"Hoja '{data['balance_sheet_name']}' no encontrada")
            
            # Leer hojas con manejo de errores
            data['balance'] = self._read_balance_sheet(xls, data['balance_sheet_name'])
            data['eri'] = self._read_eri_sheet(xls)
            data['resultado'] = self._read_resultado_sheet(xls)
            data['extracto'] = self._read_extracto_sheet(xls)
            
            self.consolidated_data = data
            self.logger.info("✅ Datos financieros cargados exitosamente")
            return data
            
        except Exception as e:
            self.logger.error(f"❌ Error cargando datos financieros: {e}")
            raise
    
    def _read_balance_sheet(self, xls: pd.ExcelFile, sheet_name: str) -> pd.DataFrame:
        """Leer hoja de balance"""
        df = pd.read_excel(xls, sheet_name=sheet_name, skiprows=4)
        
        # Asignar columnas esperadas
        expected_cols = ['Nivel', 'Código cuenta contable', 'Nombre cuenta contable', 
                        'Saldo inicial', 'Movimiento débito', 'Movimiento crédito', 'Saldo final']
        
        if len(df.columns) >= len(expected_cols):
            df.columns = expected_cols + [f'Extra_{i}' for i in range(len(df.columns) - len(expected_cols))]
        else:
            df.columns = [f'Col_{i}' for i in range(len(df.columns))]
        
        # Convertir columnas numéricas
        df['Saldo final'] = pd.to_numeric(df['Saldo final'], errors='coerce').fillna(0)
        df['Nivel'] = df['Nivel'].astype(str).fillna('')
        
        return df
    
    def _read_eri_sheet(self, xls: pd.ExcelFile) -> pd.DataFrame:
        """Leer hoja ERI"""
        df = pd.read_excel(xls, sheet_name='INFORME-ERI', skiprows=1)
        
        # Detectar columnas de meses
        eri_month_cols = [col for col in df.columns if re.match(r'^\w+ DE 2025$', str(col))]
        eri_month_cols_sorted = sorted(eri_month_cols, key=lambda x: self.month_index.get(x.split(' DE')[0], 99))
        
        if not eri_month_cols_sorted:
            # Fallback
            eri_month_cols = [col for col in df.columns[2:-1] if col in [f"{m} DE 2025" for m in self.month_order]]
            eri_month_cols_sorted = sorted(eri_month_cols, key=lambda x: self.month_index.get(x.split(' DE')[0], 99))
        
        # Reasignar columnas
        df.columns = ['COMPARATIVO MES', 'Nombre'] + list(df.columns[2:-1]) + ['OBSERVACIONES']
        df['Codigo'] = df['COMPARATIVO MES'].astype(str).fillna('')
        df['Nombre'] = df['Nombre'].astype(str).fillna('')
        df['Display Name'] = df.apply(lambda row: row['Nombre'] if row['Nombre'].strip() and row['Nombre'] != 'nan' else row['Codigo'], axis=1)
        
        # Convertir columnas de meses a numérico
        for col in eri_month_cols_sorted:
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
        
        return df
    
    def _read_resultado_sheet(self, xls: pd.ExcelFile) -> pd.DataFrame:
        """Leer hoja de estado de resultado"""
        df = pd.read_excel(xls, sheet_name='ESTADO RESULTADO', skiprows=3)
        
        # Asignar columnas
        num_columns = len(df.columns)
        if num_columns >= 2:
            df.columns = ['Descripcion'] + [f'Col_{i}' for i in range(num_columns - 1)]
        else:
            df.columns = ['Descripcion'] + [f'Col_{i}' for i in range(num_columns - 1)]
        
        return df
    
    def _read_extracto_sheet(self, xls: pd.ExcelFile) -> pd.DataFrame:
        """Leer hoja de extracto/carátula"""
        try:
            df = pd.read_excel(xls, sheet_name='CARATULA', skiprows=5)
            num_columns = len(df.columns)
            df.columns = [f'Column_{i}' for i in range(num_columns)]
            return df
        except Exception as e:
            self.logger.warning(f"No se pudo leer CARATULA: {e}")
            return pd.DataFrame()
    
    def calculate_financial_metrics(self) -> Dict:
        """Calcular métricas financieras principales"""
        if not self.consolidated_data:
            raise ValueError("No hay datos financieros cargados")
        
        data = self.consolidated_data
        current_month = data['current_month']
        
        # Obtener datos del mes actual
        current_month_col = f"{current_month} DE 2025"
        if current_month_col not in data['eri'].columns:
            current_month_col = data['eri'].columns[-1]  # Usar última columna como fallback
        
        # Calcular ingresos
        try:
            ingresos_row = data['resultado'].loc[data['resultado']['Descripcion'].str.contains('INGRESOS ORDINARIOS', case=False, na=False)]
            actual_ingresos = ingresos_row.iloc[0, 1] if not ingresos_row.empty else 0
        except:
            actual_ingresos = 0
        
        # Calcular gastos por categoría
        try:
            mask_admin = data['eri']['Codigo'].str.match(r'^51[0-9]{4,}', na=False) & (data['eri'][current_month_col] != 0)
            mask_otros = data['eri']['Codigo'].str.match(r'^53[0-9]{4,}', na=False) & (data['eri'][current_month_col] != 0)
            mask_venta = data['eri']['Codigo'].str.match(r'^61[0-9]{4,}', na=False) & (data['eri'][current_month_col] != 0)
            mask_prod = data['eri']['Codigo'].str.match(r'^(72|73)[0-9]{4,}', na=False) & (data['eri'][current_month_col] != 0)
            
            gastos_admin = data['eri'][mask_admin][current_month_col].sum()
            gastos_otros = data['eri'][mask_otros][current_month_col].sum()
            costos_venta = data['eri'][mask_venta][current_month_col].sum()
            costos_prod = data['eri'][mask_prod][current_month_col].sum()
            actual_total_gastos = gastos_admin + gastos_otros + costos_venta + costos_prod
            
            gastos_breakdown = data['eri'][mask_admin | mask_otros | mask_venta | mask_prod][['Display Name', current_month_col]].set_index('Display Name')[current_month_col].replace(0, np.nan).dropna()
        except Exception as e:
            self.logger.error(f"Error calculando gastos: {e}")
            gastos_admin = gastos_otros = costos_venta = costos_prod = actual_total_gastos = 0
            gastos_breakdown = pd.Series()
        
        # Calcular KPIs financieros
        kpis = self._calculate_kpis(data, actual_ingresos, actual_total_gastos)
        
        # Calcular presupuesto ejecutado
        ejecutado_ingresos_pct = (actual_ingresos / self.budget_data['ingresos_mensual']) * 100 if actual_ingresos > 0 else 0
        ejecutado_gastos_pct = (actual_total_gastos / self.budget_data['gastos_mensual']) * 100 if actual_total_gastos > 0 else 0
        
        metrics = {
            'current_month': current_month,
            'ingresos': actual_ingresos,
            'gastos_totales': actual_total_gastos,
            'gastos_breakdown': {
                'administrativos': gastos_admin,
                'otros': gastos_otros,
                'costos_venta': costos_venta,
                'costos_produccion': costos_prod
            },
            'presupuesto_ejecutado': {
                'ingresos_pct': ejecutado_ingresos_pct,
                'gastos_pct': ejecutado_gastos_pct
            },
            'kpis': kpis,
            'gastos_detalle': gastos_breakdown
        }
        
        self.logger.info(f"📊 Métricas calculadas para {current_month}")
        return metrics
    
    def _calculate_kpis(self, data: Dict, ingresos: float, gastos: float) -> Dict:
        """Calcular KPIs financieros"""
        try:
            balance = data['balance']
            resultado = data['resultado']
            eri = data['eri']
            current_month_col = f"{data['current_month']} DE 2025"
            
            # Activos totales (Clase 1)
            assets_class = balance[(balance['Nivel'] == 'Clase') & (balance['Código cuenta contable'] == '1')]
            total_assets = assets_class['Saldo final'].iloc[0] if not assets_class.empty else 0
            
            # Activos corrientes (Grupos 11, 12, 13, 14)
            current_groups = ['11', '12', '13', '14']
            activos_corrientes = balance[(balance['Nivel'] == 'Grupo') & balance['Código cuenta contable'].isin(current_groups)]['Saldo final'].sum()
            
            # Inventarios (Grupo 14)
            inventarios_group = balance[(balance['Nivel'] == 'Grupo') & (balance['Código cuenta contable'] == '14')]
            inventarios = inventarios_group['Saldo final'].iloc[0] if not inventarios_group.empty else 0
            
            # Pasivos corrientes (Clase 2)
            liabilities_class = balance[(balance['Nivel'] == 'Clase') & (balance['Código cuenta contable'] == '2')]
            pasivos_corrientes = abs(liabilities_class['Saldo final'].iloc[0]) if not liabilities_class.empty else 0
            
            # Patrimonio (Clase 3)
            equity_class = balance[(balance['Nivel'] == 'Clase') & (balance['Código cuenta contable'] == '3')]
            patrimonio = abs(equity_class['Saldo final'].iloc[0]) if not equity_class.empty else (total_assets - pasivos_corrientes)
            
            # Utilidad neta
            utilidad_row = resultado[resultado['Descripcion'].str.contains('RESULTADO DEL EJERCICIO', case=False, na=False)]
            utilidad = utilidad_row.iloc[0, 1] if not utilidad_row.empty else 0
            
            # Costos
            costos_row = resultado[resultado['Descripcion'].str.contains('COSTO DE VENTA', case=False, na=False)]
            costos = costos_row.iloc[0, 1] if not costos_row.empty else 0
            
            # Depreciación e intereses
            depreciacion = eri[eri['Codigo'].str.match(r'^51[0-9]{4}', na=False)][current_month_col].sum()
            intereses = eri[eri['Codigo'].str.match(r'^530520.*', na=False)][current_month_col].sum()
            
            # Calcular ratios
            current_ratio = activos_corrientes / pasivos_corrientes if pasivos_corrientes else 0
            quick_ratio = (activos_corrientes - inventarios) / pasivos_corrientes if pasivos_corrientes else 0
            margen_bruto = ((ingresos - costos) / ingresos) * 100 if ingresos else 0
            margen_neto = (utilidad / ingresos) * 100 if ingresos else 0
            roe = (utilidad / patrimonio) * 100 if patrimonio else 0
            deuda_patrimonio = pasivos_corrientes / patrimonio if patrimonio else 0
            rotacion_inventarios = costos / inventarios if inventarios else 0
            ebitda = utilidad + depreciacion + intereses
            
            return {
                'current_ratio': current_ratio,
                'quick_ratio': quick_ratio,
                'margen_bruto': margen_bruto,
                'margen_neto': margen_neto,
                'roe': roe,
                'deuda_patrimonio': deuda_patrimonio,
                'rotacion_inventarios': rotacion_inventarios,
                'ebitda': ebitda,
                'total_assets': total_assets,
                'activos_corrientes': activos_corrientes,
                'pasivos_corrientes': pasivos_corrientes,
                'patrimonio': patrimonio,
                'utilidad_neta': utilidad
            }
            
        except Exception as e:
            self.logger.error(f"Error calculando KPIs: {e}")
            return {}
    
    def generate_reports(self, metrics: Dict) -> Dict[str, str]:
        """Generar todos los reportes financieros"""
        current_month = metrics['current_month']
        month_lower = current_month.lower()
        
        reports = {}
        
        try:
            # 1. Balance consolidado
            reports['balance_consolidado'] = self._generate_consolidated_balance(month_lower)
            
            # 2. Presupuesto ejecutado
            reports['presupuesto_ejecutado'] = self._generate_budget_execution(metrics, month_lower)
            
            # 3. KPIs financieros
            reports['kpis_financieros'] = self._generate_kpi_report(metrics, month_lower)
            
            # 4. Informe para junta directiva
            reports['informe_junta'] = self._generate_board_report(metrics, month_lower)
            
            # 5. Gráficos
            reports['graficos'] = self._generate_charts(metrics, month_lower)
            
            self.logger.info("✅ Todos los reportes generados exitosamente")
            return reports
            
        except Exception as e:
            self.logger.error(f"❌ Error generando reportes: {e}")
            return {}
    
    def _generate_consolidated_balance(self, month_lower: str) -> str:
        """Generar balance consolidado"""
        try:
            # Consolidar balances de todos los meses
            balance_sheets = []
            for month in self.month_order:
                sheet_name = f'BALANCE {month}'
                if sheet_name in self.consolidated_data['sheets']:
                    df_month = pd.read_excel(self.consolidated_data['file_path'], sheet_name=sheet_name, skiprows=4)
                    df_month.columns = ['Nivel', 'Código cuenta contable', 'Nombre cuenta contable', 
                                      'Saldo inicial', 'Movimiento débito', 'Movimiento crédito', 'Saldo final'] + [f'Extra_{i}' for i in range(len(df_month.columns) - 7)]
                    df_month['Month'] = month
                    balance_sheets.append(df_month)
            
            if balance_sheets:
                df_consolidated = pd.concat(balance_sheets, ignore_index=True)
                df_consolidated = df_consolidated[df_consolidated['Nivel'] == 'Clase']
                
                file_path = os.path.join(self.reports_dir, f"balance_consolidado_{month_lower}_2025.xlsx")
                df_consolidated.to_excel(file_path, index=False)
                return file_path
        except Exception as e:
            self.logger.error(f"Error generando balance consolidado: {e}")
        return None
    
    def _generate_budget_execution(self, metrics: Dict, month_lower: str) -> str:
        """Generar reporte de presupuesto ejecutado"""
        try:
            current_month = metrics['current_month']
            ingresos = metrics['ingresos']
            gastos = metrics['gastos_totales']
            
            df_ejecutado = pd.DataFrame({
                'Categoría': ['Ingresos', 'Gastos Totales'],
                f'Actual {current_month}': [ingresos, gastos],
                'Presupuesto Mensual': [self.budget_data['ingresos_mensual'], self.budget_data['gastos_mensual']],
                '% Ejecutado': [metrics['presupuesto_ejecutado']['ingresos_pct'], metrics['presupuesto_ejecutado']['gastos_pct']]
            })
            
            file_path = os.path.join(self.reports_dir, f"presupuesto_ejecutado_{month_lower}_2025.xlsx")
            df_ejecutado.to_excel(file_path, index=False)
            return file_path
        except Exception as e:
            self.logger.error(f"Error generando presupuesto ejecutado: {e}")
        return None
    
    def _generate_kpi_report(self, metrics: Dict, month_lower: str) -> str:
        """Generar reporte de KPIs"""
        try:
            current_month = metrics['current_month']
            kpis = metrics['kpis']
            
            df_kpis = pd.DataFrame({
                'KPI': ['Current Ratio', 'Quick Ratio', 'Margen Bruto %', 'Margen Neto %', 
                       'ROE %', 'Deuda/Patrimonio', 'Rotación Inventarios', 'EBITDA'],
                f'Valor {current_month} 2025': [
                    round(kpis.get('current_ratio', 0), 2),
                    round(kpis.get('quick_ratio', 0), 2),
                    round(kpis.get('margen_bruto', 0), 2),
                    round(kpis.get('margen_neto', 0), 2),
                    round(kpis.get('roe', 0), 2),
                    round(kpis.get('deuda_patrimonio', 0), 2),
                    round(kpis.get('rotacion_inventarios', 0), 2),
                    round(kpis.get('ebitda', 0), 2)
                ]
            })
            
            file_path = os.path.join(self.reports_dir, f"kpis_financieros_{month_lower}_2025.xlsx")
            df_kpis.to_excel(file_path, index=False)
            return file_path
        except Exception as e:
            self.logger.error(f"Error generando KPIs: {e}")
        return None
    
    def _generate_board_report(self, metrics: Dict, month_lower: str) -> str:
        """Generar informe para junta directiva"""
        try:
            current_month = metrics['current_month']
            kpis = metrics['kpis']
            
            doc = Document()
            doc.add_heading(f'Informe para Junta Directiva - {current_month} 2025', 0)
            
            # Resumen financiero
            doc.add_paragraph(f'Resumen Financiero:\n- Ingresos: ${metrics["ingresos"]:,.0f}\n- Utilidad Neta: ${kpis.get("utilidad_neta", 0):,.0f}\n- Activos Totales: ${kpis.get("total_assets", 0):,.0f}\n- Pasivos Totales: ${kpis.get("pasivos_corrientes", 0):,.0f}')
            
            # Ratios financieros
            doc.add_heading('Ratios Financieros', level=1)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Ratio'
            hdr_cells[1].text = 'Valor'
            
            for kpi_name, value in [
                ('Current Ratio', kpis.get('current_ratio', 0)),
                ('Quick Ratio', kpis.get('quick_ratio', 0)),
                ('Margen Bruto %', kpis.get('margen_bruto', 0)),
                ('Margen Neto %', kpis.get('margen_neto', 0)),
                ('ROE %', kpis.get('roe', 0)),
                ('EBITDA', kpis.get('ebitda', 0))
            ]:
                row_cells = table.add_row().cells
                row_cells[0].text = kpi_name
                row_cells[1].text = f"{float(value):.2f}"
            
            # Recomendaciones
            doc.add_heading('Recomendaciones', level=1)
            recomendaciones = [
                "- Monitorear liquidez (Current Ratio bajo si < 1.5).",
                "- Revisar costos de venta elevados.",
                "- Optimizar gastos administrativos."
            ]
            for rec in recomendaciones:
                doc.add_paragraph(rec)
            
            file_path = os.path.join(self.reports_dir, f"informe_junta_{month_lower}_2025.docx")
            doc.save(file_path)
            return file_path
        except Exception as e:
            self.logger.error(f"Error generando informe junta: {e}")
        return None
    
    def _generate_charts(self, metrics: Dict, month_lower: str) -> List[str]:
        """Generar gráficos financieros"""
        chart_files = []
        
        try:
            # 1. Gráfico de gastos mensuales
            chart_files.append(self._generate_monthly_spending_chart(month_lower))
            
            # 2. Dashboard de KPIs
            chart_files.append(self._generate_kpi_dashboard(metrics, month_lower))
            
            # 3. Distribución de gastos (pie chart)
            chart_files.append(self._generate_expense_distribution_chart(metrics, month_lower))
            
            # 4. Dashboard completo
            chart_files.append(self._generate_comprehensive_dashboard(metrics, month_lower))
            
        except Exception as e:
            self.logger.error(f"Error generando gráficos: {e}")
        
        return [f for f in chart_files if f]
    
    def _generate_monthly_spending_chart(self, month_lower: str) -> str:
        """Generar gráfico de gastos mensuales"""
        try:
            data = self.consolidated_data
            eri = data['eri']
            
            # Calcular gastos por mes
            months = [col for col in eri.columns if re.match(r'^\w+ DE 2025$', str(col))]
            months_sorted = sorted(months, key=lambda x: self.month_index.get(x.split(' DE')[0], 99))
            
            total_spending = []
            for month in months_sorted:
                mask = eri['Codigo'].str.match(r'^(51|53|61|72|73)[0-9]{4,}', na=False) & (eri[month] != 0)
                total_spending.append(eri[mask][month].sum())
            
            plt.figure(figsize=(12, 6))
            plt.bar(months_sorted, total_spending, color='skyblue')
            plt.title(f'Gastos Mensuales - Enero a {data["current_month"]} 2025')
            plt.xlabel('Mes')
            plt.ylabel('Gastos Totales (COP)')
            plt.xticks(rotation=45)
            plt.grid(axis='y', linestyle='--', alpha=0.7)
            
            file_path = os.path.join(self.reports_dir, f"gastos_mensuales_{month_lower}_2025.png")
            plt.savefig(file_path, bbox_inches='tight', dpi=300)
            plt.close()
            return file_path
        except Exception as e:
            self.logger.error(f"Error generando gráfico mensual: {e}")
        return None
    
    def _generate_kpi_dashboard(self, metrics: Dict, month_lower: str) -> str:
        """Generar dashboard de KPIs"""
        try:
            current_month = metrics['current_month']
            kpis = metrics['kpis']
            
            kpi_names = ['Current Ratio', 'Quick Ratio', 'Margen Bruto %', 'Margen Neto %', 
                        'ROE %', 'Deuda/Patrimonio', 'Rotación Inventarios', 'EBITDA']
            kpi_values = [
                kpis.get('current_ratio', 0),
                kpis.get('quick_ratio', 0),
                kpis.get('margen_bruto', 0),
                kpis.get('margen_neto', 0),
                kpis.get('roe', 0),
                kpis.get('deuda_patrimonio', 0),
                kpis.get('rotacion_inventarios', 0),
                kpis.get('ebitda', 0)
            ]
            
            plt.figure(figsize=(10, 6))
            plt.bar(kpi_names, kpi_values, color='teal')
            plt.title(f'KPIs Financieros - {current_month} 2025')
            plt.xlabel('KPI')
            plt.ylabel('Valor')
            plt.xticks(rotation=45, ha='right')
            plt.grid(axis='y', linestyle='--', alpha=0.7)
            
            file_path = os.path.join(self.reports_dir, f"kpi_dashboard_{month_lower}_2025.png")
            plt.savefig(file_path, bbox_inches='tight', dpi=300)
            plt.close()
            return file_path
        except Exception as e:
            self.logger.error(f"Error generando dashboard KPIs: {e}")
        return None
    
    def _generate_expense_distribution_chart(self, metrics: Dict, month_lower: str) -> str:
        """Generar gráfico de distribución de gastos"""
        try:
            current_month = metrics['current_month']
            gastos_breakdown = metrics['gastos_breakdown']
            
            # Crear datos para el pie chart
            categories = []
            values = []
            for cat, val in gastos_breakdown.items():
                if val > 0:
                    categories.append(cat.replace('_', ' ').title())
                    values.append(val)
            
            if not values:
                return None
            
            plt.figure(figsize=(10, 8))
            colors = ['#FF9999', '#66B2FF', '#99FF99', '#FFCC99', '#FF99CC']
            wedges, texts, autotexts = plt.pie(values, labels=categories, autopct='%1.1f%%', 
                                             colors=colors[:len(values)], startangle=90)
            
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
            
            plt.title(f'Distribución de Gastos - {current_month} 2025', fontsize=14, fontweight='bold')
            plt.axis('equal')
            
            file_path = os.path.join(self.reports_dir, f"distribucion_gastos_{month_lower}_2025.png")
            plt.savefig(file_path, bbox_inches='tight', dpi=300)
            plt.close()
            return file_path
        except Exception as e:
            self.logger.error(f"Error generando distribución gastos: {e}")
        return None
    
    def _generate_comprehensive_dashboard(self, metrics: Dict, month_lower: str) -> str:
        """Generar dashboard completo"""
        try:
            current_month = metrics['current_month']
            
            fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 12))
            
            # 1. Presupuesto vs Actual
            categories = ['Ingresos', 'Gastos']
            budget_values = [self.budget_data['ingresos_mensual'], self.budget_data['gastos_mensual']]
            actual_values = [metrics['ingresos'], metrics['gastos_totales']]
            
            x = np.arange(len(categories))
            width = 0.35
            
            ax1.bar(x - width/2, budget_values, width, label='Presupuesto', color='lightgreen', alpha=0.8)
            ax1.bar(x + width/2, actual_values, width, label='Actual', color='lightcoral', alpha=0.8)
            ax1.set_xlabel('Categoría')
            ax1.set_ylabel('Valor (COP)')
            ax1.set_title('Presupuesto vs Actual')
            ax1.set_xticks(x)
            ax1.set_xticklabels(categories)
            ax1.legend()
            ax1.grid(True, alpha=0.3)
            
            # 2. KPIs principales
            kpi_names = ['Current Ratio', 'Quick Ratio', 'Margen Bruto %', 'Margen Neto %']
            kpi_values = [
                metrics['kpis'].get('current_ratio', 0),
                metrics['kpis'].get('quick_ratio', 0),
                metrics['kpis'].get('margen_bruto', 0),
                metrics['kpis'].get('margen_neto', 0)
            ]
            
            ax2.bar(kpi_names, kpi_values, color='teal', alpha=0.7)
            ax2.set_title('KPIs Principales')
            ax2.set_ylabel('Valor')
            ax2.tick_params(axis='x', rotation=45)
            ax2.grid(True, alpha=0.3)
            
            # 3. Distribución de gastos
            gastos_breakdown = metrics['gastos_breakdown']
            categories_gastos = [k.replace('_', ' ').title() for k, v in gastos_breakdown.items() if v > 0]
            values_gastos = [v for v in gastos_breakdown.values() if v > 0]
            
            if values_gastos:
                ax3.pie(values_gastos, labels=categories_gastos, autopct='%1.1f%%', startangle=90)
                ax3.set_title('Distribución de Gastos')
            
            # 4. Resumen ejecutivo
            ax4.axis('off')
            summary_text = f"""
            RESUMEN EJECUTIVO - {current_month} 2025
            
            📊 Ingresos: ${metrics['ingresos']:,.0f}
            💰 Gastos: ${metrics['gastos_totales']:,.0f}
            📈 EBITDA: ${metrics['kpis'].get('ebitda', 0):,.0f}
            🏦 Current Ratio: {metrics['kpis'].get('current_ratio', 0):.2f}
            📊 Margen Neto: {metrics['kpis'].get('margen_neto', 0):.1f}%
            """
            ax4.text(0.1, 0.5, summary_text, fontsize=12, verticalalignment='center',
                    bbox=dict(boxstyle="round,pad=0.3", facecolor="lightblue", alpha=0.5))
            
            plt.tight_layout()
            
            file_path = os.path.join(self.reports_dir, f"dashboard_completo_{month_lower}_2025.png")
            plt.savefig(file_path, bbox_inches='tight', dpi=300)
            plt.close()
            return file_path
        except Exception as e:
            self.logger.error(f"Error generando dashboard completo: {e}")
        return None
    
    def send_email_report(self, reports: Dict[str, str], metrics: Dict) -> bool:
        """Enviar reporte por email"""
        try:
            current_month = metrics['current_month']
            
            # Crear mensaje
            msg = MIMEMultipart()
            msg['From'] = self.email_config['sender_email']
            msg['To'] = self.email_config['recipient_email']
            msg['Subject'] = f'📊 Reporte CFO SuperBincent - {current_month} 2025'
            
            # Crear cuerpo del email
            html_body = f"""
            <html>
            <body>
                <h2>📊 Reporte Financiero SuperBincent - {current_month} 2025</h2>
                
                <h3>📈 Resumen Ejecutivo:</h3>
                <ul>
                    <li><strong>Ingresos:</strong> ${metrics['ingresos']:,.0f} COP</li>
                    <li><strong>Gastos Totales:</strong> ${metrics['gastos_totales']:,.0f} COP</li>
                    <li><strong>% Ejecutado Gastos:</strong> {metrics['presupuesto_ejecutado']['gastos_pct']:.1f}%</li>
                    <li><strong>Current Ratio:</strong> {metrics['kpis'].get('current_ratio', 0):.2f}</li>
                    <li><strong>EBITDA:</strong> ${metrics['kpis'].get('ebitda', 0):,.0f} COP</li>
                </ul>
                
                <h3>📎 Archivos Adjuntos:</h3>
                <ul>
                    <li>📊 Balance Consolidado (Excel)</li>
                    <li>💰 Presupuesto Ejecutado (Excel)</li>
                    <li>📈 KPIs Financieros (Excel)</li>
                    <li>📋 Informe para Junta Directiva (Word)</li>
                    <li>📊 Gráficos y Visualizaciones (PNG)</li>
                </ul>
                
                <p><em>Reporte generado automáticamente por SuperBincent CFO Bot</em></p>
            </body>
            </html>
            """
            
            msg.attach(MIMEText(html_body, 'html'))
            
            # Adjuntar archivos
            for report_type, file_path in reports.items():
                if file_path and os.path.exists(file_path):
                    with open(file_path, 'rb') as f:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(f.read())
                        encoders.encode_base64(part)
                        part.add_header('Content-Disposition', 
                                      f'attachment; filename={os.path.basename(file_path)}')
                        msg.attach(part)
            
            # Enviar email
            server = smtplib.SMTP(self.email_config['smtp_server'], self.email_config['smtp_port'])
            server.starttls()
            server.login(self.email_config['sender_email'], self.email_config['sender_password'])
            server.send_message(msg)
            server.quit()
            
            self.logger.info(f"✅ Email enviado exitosamente a {self.email_config['recipient_email']}")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Error enviando email: {e}")
            return False
    
    def update_from_entry(self, invoice_data: Dict) -> Dict:
        """Actualizar análisis financiero desde entrada de factura"""
        try:
            self.logger.info("🔄 Actualizando análisis financiero desde entrada de factura")
            
            # Detectar archivo financiero más reciente
            file_path = self.detect_latest_financial_file()
            if not file_path:
                self.logger.warning("No se encontró archivo financiero para actualizar")
                return {}
            
            # Cargar datos financieros
            financial_data = self.load_financial_data(file_path)
            
            # Calcular métricas
            metrics = self.calculate_financial_metrics()
            
            # Generar reportes
            reports = self.generate_reports(metrics)
            
            # Enviar email (opcional)
            # self.send_email_report(reports, metrics)
            
            self.logger.info("✅ Análisis financiero actualizado exitosamente")
            return {
                'metrics': metrics,
                'reports': reports,
                'status': 'success'
            }
            
        except Exception as e:
            self.logger.error(f"❌ Error actualizando análisis financiero: {e}")
            return {'status': 'error', 'error': str(e)}
    
    def run_full_analysis(self) -> Dict:
        """Ejecutar análisis financiero completo"""
        try:
            self.logger.info("🚀 Iniciando análisis financiero completo")
            
            # Detectar y cargar archivo financiero
            file_path = self.detect_latest_financial_file()
            if not file_path:
                raise ValueError("No se encontró archivo financiero")
            
            financial_data = self.load_financial_data(file_path)
            
            # Calcular métricas
            metrics = self.calculate_financial_metrics()
            
            # Generar reportes
            reports = self.generate_reports(metrics)
            
            # Enviar email
            email_sent = self.send_email_report(reports, metrics)
            
            result = {
                'status': 'success',
                'metrics': metrics,
                'reports': reports,
                'email_sent': email_sent,
                'files_generated': len([f for f in reports.values() if f and os.path.exists(f)])
            }
            
            self.logger.info(f"✅ Análisis completo finalizado - {result['files_generated']} archivos generados")
            return result
            
        except Exception as e:
            self.logger.error(f"❌ Error en análisis completo: {e}")
            return {'status': 'error', 'error': str(e)}

def main():
    """Función principal para ejecutar análisis financiero"""
    print("🚀 SuperBincent - Análisis Financiero")
    print("=" * 50)
    
    # Crear analizador
    analyzer = FinancialAnalyzer()
    
    # Ejecutar análisis completo
    result = analyzer.run_full_analysis()
    
    if result['status'] == 'success':
        print("✅ Análisis financiero completado exitosamente")
        print(f"📊 Métricas calculadas para {result['metrics']['current_month']}")
        print(f"📁 {result['files_generated']} archivos generados")
        print(f"📧 Email enviado: {'Sí' if result['email_sent'] else 'No'}")
    else:
        print(f"❌ Error en análisis: {result['error']}")

if __name__ == "__main__":
    main()