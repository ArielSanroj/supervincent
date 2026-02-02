"""
Report generation module for financial analysis.
"""

import os
import logging
from typing import Dict, Optional

import pandas as pd
from docx import Document

from .constants import DEFAULT_BUDGET

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Generates financial reports in various formats."""

    def __init__(self, reports_dir: str, budget_data: Dict = None):
        self.reports_dir = reports_dir
        self.budget_data = budget_data or DEFAULT_BUDGET
        os.makedirs(reports_dir, exist_ok=True)

    def generate_all_reports(self, metrics: Dict, data: Dict) -> Dict[str, str]:
        """Generate all financial reports."""
        current_month = metrics['current_month']
        month_lower = current_month.lower()

        reports = {}

        try:
            reports['balance_consolidado'] = self._generate_consolidated_balance(data, month_lower)
            reports['presupuesto_ejecutado'] = self._generate_budget_execution(metrics, month_lower)
            reports['kpis_financieros'] = self._generate_kpi_report(metrics, month_lower)
            reports['informe_junta'] = self._generate_board_report(metrics, month_lower)

            logger.info("All reports generated successfully")
            return reports
        except Exception as e:
            logger.error(f"Error generating reports: {e}")
            return {}

    def _generate_consolidated_balance(self, data: Dict, month_lower: str) -> Optional[str]:
        """Generate consolidated balance report."""
        try:
            from .constants import MONTH_ORDER

            balance_sheets = []
            for month in MONTH_ORDER:
                sheet_name = f'BALANCE {month}'
                if sheet_name in data['sheets']:
                    df_month = pd.read_excel(data['file_path'], sheet_name=sheet_name, skiprows=4)
                    df_month.columns = [
                        'Nivel', 'Codigo cuenta contable', 'Nombre cuenta contable',
                        'Saldo inicial', 'Movimiento debito', 'Movimiento credito', 'Saldo final'
                    ] + [f'Extra_{i}' for i in range(len(df_month.columns) - 7)]
                    df_month['Month'] = month
                    balance_sheets.append(df_month)

            if balance_sheets:
                df_consolidated = pd.concat(balance_sheets, ignore_index=True)
                df_consolidated = df_consolidated[df_consolidated['Nivel'] == 'Clase']

                file_path = os.path.join(self.reports_dir, f"balance_consolidado_{month_lower}_2025.xlsx")
                df_consolidated.to_excel(file_path, index=False)
                return file_path
        except Exception as e:
            logger.error(f"Error generating consolidated balance: {e}")
        return None

    def _generate_budget_execution(self, metrics: Dict, month_lower: str) -> Optional[str]:
        """Generate budget execution report."""
        try:
            current_month = metrics['current_month']
            ingresos = metrics['ingresos']
            gastos = metrics['gastos_totales']

            df_ejecutado = pd.DataFrame({
                'Categoria': ['Ingresos', 'Gastos Totales'],
                f'Actual {current_month}': [ingresos, gastos],
                'Presupuesto Mensual': [
                    self.budget_data['ingresos_mensual'],
                    self.budget_data['gastos_mensual']
                ],
                '% Ejecutado': [
                    metrics['presupuesto_ejecutado']['ingresos_pct'],
                    metrics['presupuesto_ejecutado']['gastos_pct']
                ]
            })

            file_path = os.path.join(self.reports_dir, f"presupuesto_ejecutado_{month_lower}_2025.xlsx")
            df_ejecutado.to_excel(file_path, index=False)
            return file_path
        except Exception as e:
            logger.error(f"Error generating budget execution: {e}")
        return None

    def _generate_kpi_report(self, metrics: Dict, month_lower: str) -> Optional[str]:
        """Generate KPI report."""
        try:
            current_month = metrics['current_month']
            kpis = metrics['kpis']

            df_kpis = pd.DataFrame({
                'KPI': [
                    'Current Ratio', 'Quick Ratio', 'Margen Bruto %', 'Margen Neto %',
                    'ROE %', 'Deuda/Patrimonio', 'Rotacion Inventarios', 'EBITDA'
                ],
                f'Valor {current_month} 2025': [
                    round(kpis.get('current_ratio', 0), 2),
                    round(kpis.get('quick_ratio', 0), 2),
                    round(kpis.get('margen_bruto', 0), 2),
                    round(kpis.get('margen_neto', 0), 2),
                    round(kpis.get('roe', 0), 2),
                    round(kpis.get('deuda_patrimonio', 0), 2),
                    round(kpis.get('rotacion_inventarios', 0), 2),
                    round(kpis.get('ebitda', 0), 2)
                ]
            })

            file_path = os.path.join(self.reports_dir, f"kpis_financieros_{month_lower}_2025.xlsx")
            df_kpis.to_excel(file_path, index=False)
            return file_path
        except Exception as e:
            logger.error(f"Error generating KPIs: {e}")
        return None

    def _generate_board_report(self, metrics: Dict, month_lower: str) -> Optional[str]:
        """Generate board report in Word format."""
        try:
            current_month = metrics['current_month']
            kpis = metrics['kpis']

            doc = Document()
            doc.add_heading(f'Informe para Junta Directiva - {current_month} 2025', 0)

            summary = (
                f'Resumen Financiero:\n'
                f'- Ingresos: ${metrics["ingresos"]:,.0f}\n'
                f'- Utilidad Neta: ${kpis.get("utilidad_neta", 0):,.0f}\n'
                f'- Activos Totales: ${kpis.get("total_assets", 0):,.0f}\n'
                f'- Pasivos Totales: ${kpis.get("pasivos_corrientes", 0):,.0f}'
            )
            doc.add_paragraph(summary)

            doc.add_heading('Ratios Financieros', level=1)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Ratio'
            hdr_cells[1].text = 'Valor'

            kpi_items = [
                ('Current Ratio', kpis.get('current_ratio', 0)),
                ('Quick Ratio', kpis.get('quick_ratio', 0)),
                ('Margen Bruto %', kpis.get('margen_bruto', 0)),
                ('Margen Neto %', kpis.get('margen_neto', 0)),
                ('ROE %', kpis.get('roe', 0)),
                ('EBITDA', kpis.get('ebitda', 0))
            ]

            for kpi_name, value in kpi_items:
                row_cells = table.add_row().cells
                row_cells[0].text = kpi_name
                row_cells[1].text = f"{float(value):.2f}"

            doc.add_heading('Recomendaciones', level=1)
            recommendations = [
                "- Monitorear liquidez (Current Ratio bajo si < 1.5).",
                "- Revisar costos de venta elevados.",
                "- Optimizar gastos administrativos."
            ]
            for rec in recommendations:
                doc.add_paragraph(rec)

            file_path = os.path.join(self.reports_dir, f"informe_junta_{month_lower}_2025.docx")
            doc.save(file_path)
            return file_path
        except Exception as e:
            logger.error(f"Error generating board report: {e}")
        return None
